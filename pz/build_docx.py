#!/usr/bin/env python3
"""
Convert pz/01_analysis.md → pz/01_analysis.docx
ITMO VKR formatting requirements (SAT_PZ_REQUIREMENTS.md):
  - A4, margins L30 R15 T20 B20 mm
  - Times New Roman 14 pt, 1.5 line spacing, justified
  - First-line indent 1.25 cm
  - Headings: bold, from indent, no period; chapter starts on new page
  - Tables: caption "Таблица N.M – Name" above left; actual table (not image)
  - Page numbers: center bottom, arabic
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = Path(__file__).parent / '01_analysis.md'
OUTPUT = Path(__file__).parent / '01_analysis.docx'

TNR       = 'Times New Roman'
COURIER   = 'Courier New'
FONT_PT   = Pt(14)
FONT_SM   = Pt(11)   # table body
INDENT    = Cm(1.25)
ONE5      = WD_LINE_SPACING.ONE_POINT_FIVE
JUSTIFY   = WD_ALIGN_PARAGRAPH.JUSTIFY
LEFT      = WD_ALIGN_PARAGRAPH.LEFT
CENTER    = WD_ALIGN_PARAGRAPH.CENTER


# ─── helpers ────────────────────────────────────────────────────────────────

def _page_num_field(para):
    run = para.add_run()
    for kind in ('begin', 'end'):
        if kind == 'begin':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); run._r.append(fc)
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'; run._r.append(it)
        else:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end'); run._r.append(fc)
    run.font.name = TNR; run.font.size = FONT_PT


def _apply_run(run, bold=False, italic=False, code=False, size=None):
    run.font.name = COURIER if code else TNR
    run.font.size = size or FONT_PT
    run.bold  = bold
    run.italic = italic


def _set_para_fmt(pf, indent=True, align=JUSTIFY, before=0, after=0, page_break=False):
    pf.first_line_indent = INDENT if indent else Pt(0)
    pf.alignment         = align
    pf.line_spacing_rule = ONE5
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.page_break_before = page_break


def _inline(para, text, default_size=None):
    """Add runs to *para* respecting **bold**, *italic*, `code`, [@cite]."""
    tok = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in tok:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2]); _apply_run(run, bold=True, size=default_size)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1]); _apply_run(run, italic=True, size=default_size)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = para.add_run(part[1:-1]); _apply_run(run, code=True, size=default_size)
        else:
            run = para.add_run(part); _apply_run(run, size=default_size)


# ─── document factory ───────────────────────────────────────────────────────

def make_doc():
    doc = Document()

    # page geometry
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(30)
    sec.right_margin  = Mm(15)
    sec.top_margin    = Mm(20)
    sec.bottom_margin = Mm(20)

    # default Normal style
    ns = doc.styles['Normal']
    ns.font.name = TNR; ns.font.size = FONT_PT
    ns.paragraph_format.line_spacing_rule = ONE5
    ns.paragraph_format.alignment         = JUSTIFY
    ns.paragraph_format.space_before      = Pt(0)
    ns.paragraph_format.space_after       = Pt(0)

    # footer: page number centred
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = CENTER
    _set_para_fmt(fp.paragraph_format, indent=False, align=CENTER)
    _page_num_field(fp)

    return doc


# ─── block adders ───────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    para = doc.add_paragraph()
    _set_para_fmt(para.paragraph_format, indent=True, align=JUSTIFY,
                  page_break=(level == 1))
    run = para.add_run(text)
    _apply_run(run, bold=True)
    return para


def add_body(doc, text, indent=True):
    para = doc.add_paragraph()
    _set_para_fmt(para.paragraph_format, indent=indent, align=JUSTIFY)
    _inline(para, text)
    return para


def add_table_caption(doc, text):
    para = doc.add_paragraph()
    _set_para_fmt(para.paragraph_format, indent=False, align=LEFT, before=6, after=3)
    _inline(para, text)
    return para


def _parse_md_table(lines):
    """Return (header_row, data_rows) as lists of string lists."""
    header, data = None, []
    for ln in lines:
        s = ln.strip()
        if re.match(r'^\|[-:\s|]+\|$', s):
            continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        if header is None:
            header = cells
        else:
            data.append(cells)
    return header or [], data


def add_md_table(doc, md_lines, usable_width_mm=165):
    """Build a Word table from markdown table lines."""
    header, data_rows = _parse_md_table(md_lines)
    if not header:
        return

    all_rows = [header] + data_rows
    ncols = max(len(r) for r in all_rows)

    # Pad short rows
    all_rows = [r + [''] * (ncols - len(r)) for r in all_rows]

    tbl = doc.add_table(rows=len(all_rows), cols=ncols)
    tbl.style = 'Table Grid'

    # Set table width to page usable width
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl_el.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(usable_width_mm * 56.7)))  # twips ≈ mm × 56.7
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    for r_idx, row_cells in enumerate(all_rows):
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row_cells):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            _set_para_fmt(para.paragraph_format, indent=False, align=LEFT,
                          before=2, after=2)
            _inline(para, cell_text, default_size=FONT_SM)
            if is_header:
                for run in para.runs:
                    run.bold = True

    return tbl


# ─── main parser ────────────────────────────────────────────────────────────

TABLE_CAP_RE = re.compile(r'^Таблица\s+\d+\.\d+\s+[–—]\s+')

def convert():
    text = INPUT.read_text(encoding='utf-8')
    lines = text.split('\n')
    doc   = make_doc()

    i = 0
    pending_cap = None  # table caption waiting for its table

    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()
        i  += 1

        if not s:
            continue

        # ── headings ──────────────────────────────────
        if s.startswith('#### '):
            add_heading(doc, s[5:], 4); continue
        if s.startswith('### '):
            add_heading(doc, s[4:], 3); continue
        if s.startswith('## '):
            add_heading(doc, s[3:], 2); continue
        if s.startswith('# '):
            add_heading(doc, s[2:], 1); continue

        # ── table caption ─────────────────────────────
        if TABLE_CAP_RE.match(s):
            pending_cap = s
            continue

        # ── table body ────────────────────────────────
        if s.startswith('|'):
            md_lines = [raw]
            while i < len(lines) and lines[i].strip().startswith('|'):
                md_lines.append(lines[i]); i += 1
            if pending_cap:
                add_table_caption(doc, pending_cap)
                pending_cap = None
            add_md_table(doc, md_lines)
            # small spacer after table
            sp = doc.add_paragraph()
            _set_para_fmt(sp.paragraph_format, indent=False)
            continue

        # ── regular paragraph ─────────────────────────
        if pending_cap:
            # orphan caption (no table followed) — emit as body text
            add_body(doc, pending_cap, indent=False)
            pending_cap = None

        add_body(doc, s)

    doc.save(str(OUTPUT))
    print(f'Saved: {OUTPUT}')


if __name__ == '__main__':
    convert()
